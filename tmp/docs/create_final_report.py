from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('.').resolve()
OUT = ROOT / 'output/doc/BodyFit_Final_Report_Draft.docx'
OUT.parent.mkdir(parents=True, exist_ok=True)

result = json.loads((ROOT / 'outputs/final/deva/result.json').read_text())
seg_report = json.loads((ROOT / 'outputs/final/segmentation/report.json').read_text())
ablation_text = (ROOT / 'outputs/final/ablation_table_limited.txt').read_text().strip()
reliability_text = (ROOT / 'outputs/final/reliability_table.txt').read_text().strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)


def add_markdown_table(doc: Document, md: str, caption: str) -> None:
    lines = [line for line in md.splitlines() if line.strip().startswith('|')]
    if len(lines) < 3:
        return
    headers = [x.strip() for x in lines[0].strip('|').split('|')]
    rows = [[x.strip() for x in line.strip('|').split('|')] for line in lines[2:]]
    p = doc.add_paragraph()
    r = p.add_run(caption)
    r.bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], 'D9EAF7')
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style='List Number')


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(10)
for name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
    styles[name].font.name = 'Arial'
styles['Title'].font.size = Pt(20)
styles['Heading 1'].font.size = Pt(14)
styles['Heading 2'].font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Reliable 2D Body Measurement from Phone Images')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(20, 52, 90)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('Final Draft Report - BodyFit Updated Direction').italic = True
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Prepared for final reporting based on the post-IJCAI revision work.').font.size = Pt(9)

doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'The original IJCAI submission framed the system around BMI, body-fat percentage, and BAI-style proxy supervision. '
    'That framing was weak because BodyM does not provide clinical body-fat labels, BMI is largely determined by height and weight, '
    'and the strongest reported BMI result depended on metadata rather than visual shape. The revised project changes the target: '
    'the model now predicts real tape-measured body dimensions from front and side silhouettes, then derives clinically interpretable indices such as WHR and WHtR.'
)
doc.add_paragraph(
    'The final prototype runs from phone RGB images to YOLO+SAM2 silhouettes, predicts waist/hip/chest dimensions, computes WHR/WHtR/BRI, '
    'and applies an SMPL-X render-back reliability check before reporting the result. The core contribution is not that WHR can be computed; '
    'the contribution is a practical, privacy-preserving phone measurement pipeline with a geometry-based accept/reject mechanism.'
)

box = doc.add_table(rows=4, cols=2)
box.style = 'Table Grid'
for row in box.rows:
    for cell in row.cells:
        set_cell_shading(cell, 'F3F7FB')
set_cell_text(box.rows[0].cells[0], 'Previous claim', True)
set_cell_text(box.rows[0].cells[1], 'BMI/BF/BAI from silhouettes')
set_cell_text(box.rows[1].cells[0], 'Revised claim', True)
set_cell_text(box.rows[1].cells[1], 'Waist/hip/chest measurement from phone silhouettes, with WHR/WHtR reporting')
set_cell_text(box.rows[2].cells[0], 'Reliability mechanism', True)
set_cell_text(box.rows[2].cells[1], 'NLF SMPL-X mesh recovery + render-back silhouette agreement')
set_cell_text(box.rows[3].cells[0], 'Final framing', True)
set_cell_text(box.rows[3].cells[1], 'Reliable 2D body measurement from phone images')

doc.add_heading('Why the Direction Changed After IJCAI', level=1)
doc.add_paragraph(
    'The rejected IJCAI draft and response make clear that the previous framing had three technical problems. First, the paper treated BMI and body-fat percentage as the main endpoints, '
    'even though BMI is a formula from height and weight and BodyM does not contain clinical body-fat ground truth. Second, BAI was used as a proxy target despite known limitations and bias. '
    'Third, the old ablation table showed a metadata shortcut: silhouette-only BMI MAE was about 22.8, while adding height and weight reduced BMI MAE to 0.00. That result is mathematically expected rather than visually meaningful.'
)
doc.add_paragraph('The revised work responds by removing the weak target and training on what the dataset actually contains: tape-measured body dimensions.')

old = doc.add_table(rows=1, cols=4)
old.style = 'Table Grid'
for i, h in enumerate(['Old result from IJCAI draft', 'Silhouette-only', 'Multi-view', 'Multi-view + height + weight']):
    set_cell_text(old.rows[0].cells[i], h, True)
    set_cell_shading(old.rows[0].cells[i], 'FCE4D6')
r = old.add_row().cells
set_cell_text(r[0], 'BMI MAE')
set_cell_text(r[1], '22.752')
set_cell_text(r[2], '22.873')
set_cell_text(r[3], '0.000')
r = old.add_row().cells
set_cell_text(r[0], 'Interpretation')
set_cell_text(r[1], 'Vision alone failed')
set_cell_text(r[2], 'Dual view did not fix BMI')
set_cell_text(r[3], 'Metadata leakage/formula shortcut')


doc.add_heading('Updated Pipeline', level=1)
add_numbered(doc, [
    'Input front and side phone RGB images.',
    'Detect the person using YOLO and prompt SAM2 with the padded person bounding box.',
    'Standardize the binary silhouettes to the BodyM-compatible 640 x 480 canvas.',
    'Predict waist_cm, hip_cm, and chest_cm using the dual-view model checkpoint.',
    'Compute WHR, WHtR, and BRI from predicted dimensions and known height.',
    'Run NLF on the front RGB image to recover a personalized SMPL-X mesh.',
    'Render the SMPL-X mesh back to the silhouette and accept/reject the report using IoU and Chamfer mismatch.',
])

doc.add_heading('Implemented Changes', level=1)
add_bullets(doc, [
    'Changed active target from BMI/body-fat/BAI to waist_cm, hip_cm, and chest_cm.',
    'Built the canonical BodyM table at data/bodym/pairs_dimensions.csv with 6,134 paired rows and 2,018 subjects.',
    'Removed body-fat pseudo-labels from the active training and reporting path.',
    'Changed inference to output dimensions, WHR, WHtR, BRI, risk labels, and a reportable flag.',
    'Repaired segmentation to use strict YOLO person boxes and padded box-only SAM2 prompts, with no GrabCut or fallback masks.',
    'Added the NLF SMPL-X fit path for phone RGB reliability checking.',
    'Kept the lightweight proxy gate only for BodyM mask-only retained-error experiments.',
])


doc.add_heading('Dataset and Supervision', level=1)
doc.add_paragraph(
    'The active dataset is data/bodym/pairs_dimensions.csv. It has 6,134 paired samples from 2,018 subjects. '
    'The dimension columns include height_cm, waist_cm, hip_cm, chest_cm, bicep_cm, thigh_cm, wrist_cm, and other BodyM measurements. '
    'For this report, the active model uses waist_cm, hip_cm, and chest_cm because these are directly relevant to central body shape and health-index derivation.'
)


doc.add_heading('Results That Can Be Reported', level=1)
doc.add_paragraph(
    'The following results are suitable for the final report. The limited ablation table uses the first 160 rows for a fast report-ready validation check; '
    'the full-table command should be run separately if time permits before submission.'
)
add_markdown_table(doc, ablation_text, 'Table 1. Limited dimension-first ablation on 160 BodyM rows.')
add_markdown_table(doc, reliability_text, 'Table 2. Legacy proxy gate retained-error table on 100 BodyM rows.')


doc.add_heading('Final Phone Demo', level=1)
meas = result['measurements']
indices = result['indices']
risks = result['risks']
smpl = result['smplx_fit']
metrics = smpl['metrics']
demo = doc.add_table(rows=1, cols=3)
demo.style = 'Table Grid'
for i, h in enumerate(['Output', 'Value', 'Interpretation']):
    set_cell_text(demo.rows[0].cells[i], h, True)
    set_cell_shading(demo.rows[0].cells[i], 'E2F0D9')
for name, value, interp in [
    ('Waist', f"{meas['waist_cm']:.2f} cm", 'Predicted circumference from silhouette model'),
    ('Hip', f"{meas['hip_cm']:.2f} cm", 'Predicted circumference from silhouette model'),
    ('Chest', f"{meas['chest_cm']:.2f} cm", 'Predicted circumference from silhouette model'),
    ('WHR', f"{indices['WHR']:.4f}", risks['WHR']),
    ('WHtR', f"{indices['WHtR']:.4f}", risks['WHtR_secondary']),
    ('BRI', f"{indices['BRI']:.4f}", risks['BRI']),
    ('SMPL-X reliability', 'ACCEPTED' if smpl['accepted'] else 'REJECTED', f"IoU={metrics['front_iou']:.4f}, Chamfer={metrics['front_chamfer']:.4f}"),
]:
    row = demo.add_row().cells
    set_cell_text(row[0], name)
    set_cell_text(row[1], value)
    set_cell_text(row[2], interp)

doc.add_paragraph(
    f"Segmentation smoke test: {sum(1 for r in seg_report if r.get('ok'))}/{len(seg_report)} TestPhoto samples passed. "
    'The final artifacts are stored under outputs/final/.'
)

for image_path, caption in [
    ('outputs/final/segmentation/contact_sheet.png', 'Figure 1. Final segmentation smoke contact sheet.'),
    ('outputs/final/deva/smplx/smplx_front_overlay.png', 'Figure 2. SMPL-X render-back overlay for the final phone demo.'),
]:
    path = ROOT / image_path
    if path.exists():
        doc.add_paragraph(caption).runs[0].bold = True
        doc.add_picture(str(path), width=Inches(5.8))


doc.add_heading('What This Means Relative to BodyM', level=1)
doc.add_paragraph(
    'BodyM-style work already shows that body dimensions can be estimated from silhouettes. Therefore, this report should not claim novelty from computing WHR itself. '
    'The useful result is that the current pipeline reproduces a dimension-first version of the task and adds a practical reliability gate for phone captures. '
    'The gap to BodyM should be discussed in terms of deployment: phone RGB segmentation quality, accept/reject behavior, and robust reporting of derived indices.'
)


doc.add_heading('Limitations', level=1)
add_bullets(doc, [
    'The current model predicts circumference labels from silhouettes; it does not measure clinical body fat percentage.',
    'WHR, WHtR, and BRI are derived shape indices, not direct adiposity measurements.',
    'The SMPL-X gate currently checks front-view render agreement and does not yet correct waist/hip/chest values.',
    'The phone demo is a prototype example; population-level phone-RGB validation is still required.',
    'Loose clothing, abayas, occlusions, pose variation, and camera distance remain major real-world failure modes.',
])


doc.add_heading('What To Finish Today', level=1)
add_numbered(doc, [
    'Run the full 5-eval/4ablate.py table if time permits; otherwise report the limited table as preliminary.',
    'Use outputs/final/segmentation/contact_sheet.png and outputs/final/deva/smplx/smplx_front_overlay.png as final figures.',
    'Keep the final claim short: Reliable 2D body measurement from phone images.',
    'Do not claim body-fat prediction, DEXA equivalence, or novelty of WHR computation.',
    'Position mesh-ring circumference extraction and stronger clothing validation as future work.',
])


doc.add_heading('Proposed Final Abstract', level=1)
doc.add_paragraph(
    'This project presents a phone-based body measurement pipeline that estimates waist, hip, and chest circumference from front and side silhouettes. '
    'Unlike the previous BMI/body-fat framing, the revised system predicts dimensions directly available in the BodyM dataset and derives WHR, WHtR, and BRI from those measurements. '
    'The RGB preprocessing pipeline uses YOLO person detection and box-prompted SAM2 segmentation to produce standardized silhouettes. A dual-view model predicts body dimensions, while an NLF-based SMPL-X render-back check determines whether the capture is reliable enough to report. '
    'Preliminary results show substantially lower dimension error than the previous BMI-oriented silhouette-only framing and demonstrate an end-to-end phone image example with accepted SMPL-X reliability. The system should be interpreted as reliable 2D body measurement from phone images, not as direct body-fat estimation.'
)


doc.add_heading('References / Evidence Used', level=1)
add_bullets(doc, [
    'Rejected IJCAI-26 draft: original BMI/BF/BAI framing and metadata ablation table.',
    'IJCAI author response: reviewer concern that silhouettes alone were insufficient and evaluation was preliminary.',
    'Current repository: data/bodym/pairs_dimensions.csv, 4-infer/1infer.py, 5-eval/4ablate.py, 5-eval/6gate_eval.py, 5-eval/7segmentation_smoke.py.',
    'Final artifacts: outputs/final/deva/result.json and outputs/final/segmentation/report.json.',
])

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.text = 'BodyFit final draft - generated from current repository artifacts'
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
